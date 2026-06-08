from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_border(cell):
    """為 cell 加上細框線。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def build_docx(table_json: dict, output_path: str) -> None:
    rows_data = table_json.get("rows", [])
    if not rows_data:
        doc = Document()
        doc.add_paragraph("（無法識別表格內容）")
        doc.save(output_path)
        return

    # 計算總列數與最大欄數
    total_rows = len(rows_data)
    max_cols = 0
    for row in rows_data:
        col_count = sum(c.get("colspan", 1) for c in row.get("cells", []))
        if col_count > max_cols:
            max_cols = col_count

    if max_cols == 0:
        doc = Document()
        doc.add_paragraph("（無法識別表格內容）")
        doc.save(output_path)
        return

    doc = Document()
    table = doc.add_table(rows=total_rows, cols=max_cols)

    # occupied[r][c] = True 表示該格已被合併佔用，跳過
    occupied = [[False] * max_cols for _ in range(total_rows)]

    for r_idx, row in enumerate(rows_data):
        c_cursor = 0
        for cell_data in row.get("cells", []):
            # 找下一個未被佔用的欄位
            while c_cursor < max_cols and occupied[r_idx][c_cursor]:
                c_cursor += 1
            if c_cursor >= max_cols:
                break

            text = cell_data.get("text", "")
            rowspan = max(1, int(cell_data.get("rowspan", 1)))
            colspan = max(1, int(cell_data.get("colspan", 1)))

            # 標記被佔用的區域
            for dr in range(rowspan):
                for dc in range(colspan):
                    nr, nc = r_idx + dr, c_cursor + dc
                    if nr < total_rows and nc < max_cols:
                        occupied[nr][nc] = True

            cell = table.cell(r_idx, c_cursor)
            cell.text = text

            # 合併儲存格
            if rowspan > 1 or colspan > 1:
                end_r = min(r_idx + rowspan - 1, total_rows - 1)
                end_c = min(c_cursor + colspan - 1, max_cols - 1)
                cell.merge(table.cell(end_r, end_c))
                # merge 後文字留在左上角 cell，重新設定
                table.cell(r_idx, c_cursor).text = text

            _set_cell_border(table.cell(r_idx, c_cursor))
            c_cursor += colspan

    doc.save(output_path)
